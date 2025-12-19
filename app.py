import io
import re
import time
import base64
from typing import List, Tuple, Optional, Dict

import streamlit as st
from openai import OpenAI
import fitz  # PyMuPDF
from PIL import Image

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import streamlit.components.v1 as components


# =========================
# Defaults
# =========================

DEFAULT_BASE_URL = "https://api.sambanova.ai/v1"
DEFAULT_VISION_MODEL = "Llama-4-Maverick-17B-128E-Instruct"
DEFAULT_TEXT_MODEL = "Meta-Llama-3.3-70B-Instruct"

SYSTEM_PROMPT = """Bạn là trợ lý chuyển đổi tài liệu Toán sang văn bản gõ lại.
YÊU CẦU NGHIÊM NGẶT:
1) Mọi công thức/toán học phải nằm trong dấu $...$ (inline math).
2) TUYỆT ĐỐI KHÔNG có ký tự xuống dòng bên trong $...$.
3) Không tự ý sắp xếp lại thứ tự, không đổi số câu, không gộp/tách câu.
4) Giữ xuống dòng lời giải hợp lí (giống bố cục gốc), nhưng không đưa tab \\t.
5) Nếu không chắc một ký hiệu/toán tử, hãy giữ nguyên như nhìn thấy.
6) Đầu ra chỉ là NỘI DUNG (plain text), không thêm tiêu đề/giải thích ngoài lề.
"""

VISION_USER_INSTRUCTION = """Hãy đọc chính xác nội dung trong ảnh và gõ lại.
- Giữ nguyên thứ tự dòng/ý/câu.
- Với mọi biểu thức toán học: bọc vào $...$ và đảm bảo không có xuống dòng trong $...$.
- Không dùng \\(\\), \\[\\], $$...$$; chỉ dùng $...$.
- Không dùng tab.
Trả về đúng nội dung đã gõ lại (plain text)."""

TEXT_CLEANUP_INSTRUCTION = """Bạn hãy chuẩn hóa lại văn bản sau cho đúng yêu cầu:
- Mọi công thức/toán học phải nằm trong $...$.
- Không có xuống dòng trong $...$.
- Không thêm/bớt ý, không đổi thứ tự.
- Không dùng tab.
Chỉ trả về văn bản đã chuẩn hóa."""


# =========================
# Helpers
# =========================

def make_client(api_key: str, base_url: str) -> OpenAI:
    return OpenAI(api_key=api_key, base_url=base_url)


def encode_image_bytes(img_bytes: bytes, mime: str) -> str:
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    return f"data:{mime};base64,{b64}"


def strip_tabs(text: str) -> str:
    return text.replace("\t", " ").replace("\u000b", " ")


def collapse_newlines_inside_dollars(text: str) -> str:
    def _fix_block(m: re.Match) -> str:
        inner = m.group(1)
        inner = inner.replace("\r", " ").replace("\n", " ")
        inner = re.sub(r"\s{2,}", " ", inner).strip()
        return f"${inner}$"
    return re.sub(r"\$(.*?)\$", _fix_block, text, flags=re.DOTALL)


def final_sanitize(text: str) -> str:
    text = strip_tabs(text)
    text = collapse_newlines_inside_dollars(text)
    text = re.sub(r"[ \u00A0]{3,}", "  ", text)
    return text.strip()


def pil_to_png_bytes(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def pdf_page_count(pdf_bytes: bytes) -> int:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n = len(doc)
    doc.close()
    return n


def render_pdf_page(pdf_bytes: bytes, page_index: int, dpi: int) -> Image.Image:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_index]
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
    doc.close()
    return img


def call_vision_transcribe(
    client: OpenAI,
    model: str,
    image_png_bytes: bytes,
    max_tokens: int,
    temperature: float,
) -> str:
    data_url = encode_image_bytes(image_png_bytes, "image/png")
    resp = client.chat.completions.create(
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": VISION_USER_INSTRUCTION},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ],
    )
    out = resp.choices[0].message.content or ""
    return final_sanitize(out)


def call_text_cleanup(client: OpenAI, model: str, raw_text: str, max_tokens: int) -> str:
    resp = client.chat.completions.create(
        model=model,
        temperature=0.1,
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": TEXT_CLEANUP_INSTRUCTION + "\n\n---\n\n" + raw_text},
        ],
    )
    out = resp.choices[0].message.content or ""
    return final_sanitize(out)


def transcribe_with_retry(
    client: OpenAI,
    vision_model: str,
    img: Image.Image,
    *,
    max_tokens: int,
    temperature: float,
    retries: int = 2,
    min_chars_ok: int = 40,
) -> Tuple[str, Optional[str]]:
    """
    Returns (text, error_message). error_message None if ok.
    Retry if empty/too short.
    """
    png = pil_to_png_bytes(img)
    last_err = None
    for attempt in range(retries + 1):
        try:
            txt = call_vision_transcribe(client, vision_model, png, max_tokens=max_tokens, temperature=temperature)
            if len(txt.strip()) >= min_chars_ok:
                return txt, None
            last_err = f"Kết quả quá ngắn/rỗng (len={len(txt.strip())})."
        except Exception as e:
            last_err = f"Lỗi gọi vision: {e}"
        time.sleep(0.6)
    return "", last_err


def build_docx(sections: List[Tuple[str, str]]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for idx, (title, content) in enumerate(sections, start=1):
        if title:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        lines = content.splitlines() if content else []
        if not lines:
            doc.add_paragraph("")
        else:
            for line in lines:
                doc.add_paragraph(line)

        if idx != len(sections):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================
# Paste Image (Ctrl+V) component
# =========================

def paste_image_component(key: str = "paste_img") -> Optional[bytes]:
    """
    Returns image bytes (PNG) if user pasted an image, else None.
    Works via Streamlit postMessage protocol.
    """
    html = f"""
    <div style="border:1px dashed #999;padding:12px;border-radius:10px;">
      <div style="font-size:14px;margin-bottom:6px;">
        <b>Dán ảnh tại đây (Ctrl+V)</b> — chỉ nhận ảnh từ clipboard.
      </div>
      <textarea id="ta" placeholder="Click vào đây rồi Ctrl+V..." 
        style="width:100%;height:110px;resize:vertical;font-size:14px;padding:10px;"></textarea>
      <div id="status" style="margin-top:8px;color:#555;font-size:13px;"></div>
    </div>

    <script>
      const ta = document.getElementById("ta");
      const status = document.getElementById("status");

      function sendValue(value) {{
        const msg = {{
          isStreamlitMessage: true,
          type: "streamlit:setComponentValue",
          value: value
        }};
        window.parent.postMessage(msg, "*");
      }}

      ta.addEventListener("paste", async (e) => {{
        try {{
          const items = (e.clipboardData || window.clipboardData).items;
          if (!items) return;

          for (let i = 0; i < items.length; i++) {{
            const it = items[i];
            if (it.type && it.type.startsWith("image/")) {{
              const file = it.getAsFile();
              const reader = new FileReader();
              reader.onload = () => {{
                const dataUrl = reader.result; // data:image/png;base64,...
                status.textContent = "✅ Đã nhận ảnh từ clipboard.";
                // Send base64 only to streamlit
                sendValue(dataUrl);
              }};
              reader.readAsDataURL(file);
              e.preventDefault();
              return;
            }}
          }}
          status.textContent = "⚠️ Clipboard không có ảnh.";
        }} catch(err) {{
          status.textContent = "❌ Lỗi khi đọc clipboard: " + err;
        }}
      }});
    </script>
    """
    data_url = components.html(html, height=190, key=key)
    if not data_url or not isinstance(data_url, str):
        return None
    if not data_url.startswith("data:image/"):
        return None
    # decode base64
    try:
        header, b64 = data_url.split(",", 1)
        return base64.b64decode(b64)
    except Exception:
        return None


# =========================
# Streamlit UI
# =========================

st.set_page_config(page_title="Ảnh/PDF → Word (SambaNova)", layout="wide")
st.title("📄 Ảnh / PDF → Word (.docx) bằng SambaNova")
st.caption("Nghiêm ngặt: công thức toán trong $...$ và không xuống dòng bên trong $...$.")

with st.sidebar:
    st.header("Cấu hình API")
    api_key = st.text_input("SambaNova API Key", type="password", placeholder="Nhập key của bạn…")
    base_url = st.text_input("Base URL", value=DEFAULT_BASE_URL)
    vision_model = st.text_input("Vision model", value=DEFAULT_VISION_MODEL)
    text_model = st.text_input("Text model (cleanup)", value=DEFAULT_TEXT_MODEL)

    st.divider()
    st.subheader("Chất lượng đọc PDF")
    dpi_main = st.slider("DPI chính", 120, 320, 240, 10)
    dpi_fallback = st.slider("DPI fallback (nếu trang lỗi/rỗng)", 120, 320, 180, 10)

    st.divider()
    st.subheader("Giới hạn trả lời")
    vision_max_tokens = st.slider("Vision max_tokens / trang", 1500, 8000, 6000, 250)
    cleanup_max_tokens = st.slider("Cleanup max_tokens / trang", 1500, 8000, 4000, 250)
    temperature = st.slider("temperature", 0.0, 0.8, 0.2, 0.05)

    st.divider()
    do_cleanup = st.toggle("Chuẩn hoá lại (khuyến nghị)", value=True)
    min_chars_ok = st.slider("Ngưỡng tối thiểu ký tự để coi là OK", 10, 200, 40, 5)
    retries = st.slider("Số lần retry nếu trang rỗng", 0, 4, 2, 1)

tabs = st.tabs(["📎 Tải file (PDF/Ảnh)", "📋 Dán ảnh (Ctrl+V)"])

uploads = []
pasted_images: List[Tuple[str, bytes]] = []

with tabs[0]:
    st.subheader("Tải tệp")
    uploads = st.file_uploader(
        "Chọn 1 hoặc nhiều tệp (PDF/PNG/JPG/JPEG)",
        type=["pdf", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
    )

with tabs[1]:
    st.subheader("Dán ảnh từ clipboard")
    img_bytes = paste_image_component(key="paste_1")
    if img_bytes:
        pasted_images.append(("pasted_image_1.png", img_bytes))
        st.image(img_bytes, caption="Ảnh vừa dán", use_column_width=True)
    st.caption("Mẹo: dùng Snipping Tool / PrtSc để copy ảnh, sau đó click vào ô và Ctrl+V.")

have_inputs = (uploads and len(uploads) > 0) or (len(pasted_images) > 0)

if st.button("🚀 Chuyển sang Word", type="primary", disabled=(not have_inputs or not api_key)):
    client = make_client(api_key, base_url)

    sections: List[Tuple[str, str]] = []
    report_rows: List[Dict[str, str]] = []

    total_jobs = 0
    if uploads:
        for up in uploads:
            if up.name.lower().endswith(".pdf"):
                total_jobs += max(1, pdf_page_count(up.read()))
                up.seek(0)
            else:
                total_jobs += 1
    total_jobs += len(pasted_images)

    progress = st.progress(0)
    done = 0

    # -------- Handle uploads --------
    if uploads:
        for up in uploads:
            filename = up.name
            data = up.read()

            if filename.lower().endswith(".pdf"):
                st.write(f"### 📎 PDF: {filename}")
                n_pages = pdf_page_count(data)
                st.write(f"- Số trang PDF: **{n_pages}**")

                page_texts: List[str] = []
                for pi in range(n_pages):
                    page_no = pi + 1
                    with st.spinner(f"Đang đọc {filename} — trang {page_no}/{n_pages} (DPI {dpi_main}) …"):
                        try:
                            img = render_pdf_page(data, pi, dpi=dpi_main)
                        except Exception as e:
                            # render fail -> fallback dpi
                            try:
                                img = render_pdf_page(data, pi, dpi=dpi_fallback)
                            except Exception as e2:
                                report_rows.append({
                                    "File": filename,
                                    "Trang": str(page_no),
                                    "Trạng thái": "❌ Render lỗi",
                                    "Ghi chú": f"{e} | fallback: {e2}"
                                })
                                page_texts.append("")  # giữ chỗ để không “tụt trang”
                                done += 1
                                progress.progress(min(1.0, done / max(1, total_jobs)))
                                continue

                        txt, err = transcribe_with_retry(
                            client,
                            vision_model,
                            img,
                            max_tokens=vision_max_tokens,
                            temperature=temperature,
                            retries=retries,
                            min_chars_ok=min_chars_ok,
                        )

                        if (not txt.strip()) and err:
                            # thử fallback DPI nếu DPI chính rỗng
                            if dpi_fallback != dpi_main:
                                with st.spinner(f"Trang {page_no} rỗng → thử lại DPI {dpi_fallback} …"):
                                    try:
                                        img2 = render_pdf_page(data, pi, dpi=dpi_fallback)
                                        txt2, err2 = transcribe_with_retry(
                                            client, vision_model, img2,
                                            max_tokens=vision_max_tokens,
                                            temperature=temperature,
                                            retries=retries,
                                            min_chars_ok=min_chars_ok,
                                        )
                                        if txt2.strip():
                                            txt, err = txt2, None
                                        else:
                                            err = err2 or err
                                    except Exception as e3:
                                        err = f"{err} | fallback render error: {e3}"

                        if do_cleanup and txt.strip():
                            with st.spinner(f"Chuẩn hoá trang {page_no} …"):
                                try:
                                    txt = call_text_cleanup(client, text_model, txt, max_tokens=cleanup_max_tokens)
                                except Exception as e:
                                    report_rows.append({
                                        "File": filename,
                                        "Trang": str(page_no),
                                        "Trạng thái": "⚠️ Cleanup lỗi",
                                        "Ghi chú": str(e)
                                    })

                        status = "✅ OK" if txt.strip() else "⚠️ Rỗng"
                        note = "" if txt.strip() else (err or "Không rõ lý do")
                        report_rows.append({
                            "File": filename,
                            "Trang": str(page_no),
                            "Trạng thái": status,
                            "Ghi chú": note
                        })

                        # Giữ chỗ: nếu rỗng vẫn append "" để không mất trang
                        page_texts.append(txt.strip())

                    done += 1
                    progress.progress(min(1.0, done / max(1, total_jobs)))

                # ghép theo trang (có phân cách rõ)
                merged_pages = []
                for i, t in enumerate(page_texts, start=1):
                    merged_pages.append(f"[Trang {i}]\n{t}".strip())
                merged = "\n\n".join(merged_pages).strip()

                sections.append((filename, merged if merged else ""))

            else:
                st.write(f"### 🖼️ Ảnh: {filename}")
                try:
                    img = Image.open(io.BytesIO(data)).convert("RGB")
                except Exception as e:
                    report_rows.append({"File": filename, "Trang": "-", "Trạng thái": "❌ Ảnh lỗi", "Ghi chú": str(e)})
                    continue

                with st.spinner("Đang đọc ảnh…"):
                    txt, err = transcribe_with_retry(
                        client, vision_model, img,
                        max_tokens=vision_max_tokens,
                        temperature=temperature,
                        retries=retries,
                        min_chars_ok=min_chars_ok,
                    )
                    if do_cleanup and txt.strip():
                        try:
                            txt = call_text_cleanup(client, text_model, txt, max_tokens=cleanup_max_tokens)
                        except Exception as e:
                            report_rows.append({"File": filename, "Trang": "-", "Trạng thái": "⚠️ Cleanup lỗi", "Ghi chú": str(e)})

                report_rows.append({
                    "File": filename,
                    "Trang": "-",
                    "Trạng thái": "✅ OK" if txt.strip() else "⚠️ Rỗng",
                    "Ghi chú": "" if txt.strip() else (err or "Không rõ lý do")
                })
                sections.append((filename, txt))

                done += 1
                progress.progress(min(1.0, done / max(1, total_jobs)))

    # -------- Handle pasted images --------
    for name, b in pasted_images:
        st.write(f"### 📋 Ảnh dán: {name}")
        img = Image.open(io.BytesIO(b)).convert("RGB")
        with st.spinner("Đang đọc ảnh dán…"):
            txt, err = transcribe_with_retry(
                client, vision_model, img,
                max_tokens=vision_max_tokens,
                temperature=temperature,
                retries=retries,
                min_chars_ok=min_chars_ok,
            )
            if do_cleanup and txt.strip():
                try:
                    txt = call_text_cleanup(client, text_model, txt, max_tokens=cleanup_max_tokens)
                except Exception as e:
                    report_rows.append({"File": name, "Trang": "-", "Trạng thái": "⚠️ Cleanup lỗi", "Ghi chú": str(e)})

        report_rows.append({
            "File": name,
            "Trang": "-",
            "Trạng thái": "✅ OK" if txt.strip() else "⚠️ Rỗng",
            "Ghi chú": "" if txt.strip() else (err or "Không rõ lý do")
        })
        sections.append((name, txt))

        done += 1
        progress.progress(min(1.0, done / max(1, total_jobs)))

    # Build Word
    with st.spinner("Đang tạo Word…"):
        docx_bytes = build_docx(sections)

    st.success("Xong! Tải Word bên dưới. (Có báo cáo trang nào rỗng/lỗi để thầy kiểm tra.)")

    st.download_button(
        "⬇️ Tải Word (.docx)",
        data=docx_bytes,
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.subheader("📋 Báo cáo đọc trang")
    # hiển thị report
    if report_rows:
        st.dataframe(report_rows, use_container_width=True)

else:
    if not api_key:
        st.info("Nhập SambaNova API Key ở sidebar.")
    elif not have_inputs:
        st.info("Tải PDF/ảnh hoặc dán ảnh (Ctrl+V) để bắt đầu.")
